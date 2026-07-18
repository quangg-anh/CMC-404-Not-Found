"""Extract raw text from uploaded legal files (PDF / DOC / DOCX / TXT / HTML).

Used by the ingest pipeline to turn a binary file (fetched from MinIO) into plain text that the
``LegalParser`` can split into Điều/Khoản.

Design goals:
  - Prefer extractors that preserve Vietnamese Unicode correctly (PyMuPDF before pdfplumber).
  - Repair common mojibake / bad encodings after extraction.
  - Normalize Unicode (NFC), strip control / private-use / replacement chars.
  - Fall back to OCR only when the text layer is missing or looks garbled.
  - Support legacy ``.doc`` via Word COM (Windows) or LibreOffice when available.
"""
from __future__ import annotations

import io
import logging
import os
import re
import shutil
import subprocess
import tempfile
import unicodedata
from pathlib import Path

logger = logging.getLogger(__name__)

# Vietnamese + Latin letters used to score text quality.
_LETTER_RE = re.compile(r"[A-Za-zÀ-ỹĐđ]")
# Characters that almost never appear in clean Vietnamese legal prose.
_WEIRD_RE = re.compile(r"[^\w\sÀ-ỹĐđ.,;:!?()/%+\-–—\"'“”‘’\[\]{}|/\\@#&*=<>°‰…]")
# Private-use / replacement / C0 controls except \t \n \r.
_BAD_CHARS_RE = re.compile(r"[\u0000-\u0008\u000B\u000C\u000E-\u001F\u007F-\u009F\uFFFD\uE000-\uF8FF]")


def extract_text(data: bytes, filename: str = "", mime: str = "") -> str:
    """Dispatch to the right extractor based on filename extension / MIME type."""
    if not data:
        return ""
    name = (filename or "").lower()
    m = (mime or "").lower()

    if name.endswith(".pdf") or "pdf" in m:
        raw = _from_pdf(data)
    elif name.endswith(".docx") or "wordprocessingml" in m or "opendocument.text" in m:
        raw = _from_docx(data)
    elif name.endswith(".doc") or m in {"application/msword", "application/x-msword"}:
        raw = _from_doc(data, filename or "document.doc")
    elif name.endswith((".html", ".htm")) or "html" in m:
        raw = _strip_html(_decode(data))
    elif name.endswith(".rtf") or "rtf" in m:
        raw = _from_rtf(data)
    else:
        raw = _decode(data)

    return clean_text(raw)


# ---------------------------------------------------------------------------
# Public cleaning / quality helpers (also used by tests)
# ---------------------------------------------------------------------------

def clean_text(text: str) -> str:
    """Normalize and repair extracted text so Vietnamese characters stay readable."""
    if not text:
        return ""

    text = _repair_mojibake(text)
    # Prefer composed Vietnamese (NFC): "ế" as one codepoint, not e + combining marks.
    text = unicodedata.normalize("NFC", text)
    text = _BAD_CHARS_RE.sub("", text)

    # Soft hyphen / zero-width / BOM leftovers that break word boundaries.
    for ch in ("\u00ad", "\u200b", "\u200c", "\u200d", "\ufeff", "\u2060"):
        text = text.replace(ch, "")

    text = _repair_vietnamese_ocr(text)

    # Normalize newlines and collapse runaway whitespace without destroying paragraph breaks.
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t\f\v]+", " ", text)
    text = re.sub(r" *\n *", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    # Drop empty table-cell leftovers like " | | |"
    text = re.sub(r"(?:\s*\|\s*){2,}", " ", text)
    text = re.sub(r"\s+\|\s*$", "", text, flags=re.MULTILINE)
    return text.strip()


# Common Vietnamese OCR confusions in legal PDFs (Tesseract vie on scanned decrees).
# Order matters: longer / more specific phrases first.
_OCR_PHRASE_FIXES: list[tuple[str, str]] = [
    (r"công\s*bồ\s*MƠng", "công bố lượng"),
    (r"công\s*bồ", "công bố"),
    (r"không\s*đỄtu\s*giá", "không đấu giá"),
    (r"đỄtu\s*giá", "đấu giá"),
    (r"trúng\s*đâu\s*giá", "trúng đấu giá"),
    (r"không\s*hệt", "không hết"),
    (r"Hội\s*đông\s*đâu\s*giá", "Hội đồng đấu giá"),
    (r"Hội\s*đông", "Hội đồng"),
    (r"đâu\s*giá\s*đôi\s*với", "đấu giá đối với"),
    (r"đâu\s*giá", "đấu giá"),
    (r"hạn\s*ngạch\s*thuê\s*quan", "hạn ngạch thuế quan"),
    (r"thuê\s*quan", "thuế quan"),
    (r"diễn\s*biên", "diễn biến"),
    (r"thời\s*điêm", "thời điểm"),
    (r"kêt\s*thúc", "kết thúc"),
    (r"quyêt\s*định", "quyết định"),
    (r"tiêp\s*tục", "tiếp tục"),
    (r"đề\s*Bộ\s*Công\s*Thương", "để Bộ Công Thương"),
    (r"mặt\s*hàng\s*đương", "mặt hàng đường"),
    (r"nhập\s*khâu", "nhập khẩu"),
    (r"thue\s*quan", "thuế quan"),
    (r"Trách\s*Mnem", "Trách nhiệm"),
    (r"PIọn\s*đồng", "Hội đồng"),
    (r"hmn\s*ngạch", "hạn ngạch"),
    (r"ủ_m\s*ngạch", "hạn ngạch"),
    (r"N1an\s*giao", "phân giao"),
    (r"\bQuyen\b", "Quyền"),
    (r"fflrơng", "đường"),
]


def _repair_vietnamese_ocr(text: str) -> str:
    """Fix frequent Vietnamese OCR typos without changing correct text aggressively."""
    if not text:
        return text
    out = text
    for pattern, repl in _OCR_PHRASE_FIXES:
        out = re.sub(pattern, repl, out, flags=re.IGNORECASE)
    return out


def text_quality(text: str) -> float:
    """Score in [0, 1]: higher = more likely readable Vietnamese legal prose."""
    txt = (text or "").strip()
    if len(txt) < 20:
        return 0.0
    letters = len(_LETTER_RE.findall(txt))
    weird = len(_WEIRD_RE.findall(txt))
    replacement = txt.count("\ufffd")
    letter_ratio = letters / max(len(txt), 1)
    weird_ratio = weird / max(len(txt), 1)
    score = letter_ratio - (weird_ratio * 2.0) - (replacement * 0.05)
    # Penalize dense dotted form lines / table OCR garbage.
    if re.search(r"\.{8,}|_{8,}|-{8,}", txt):
        score -= 0.15
    if txt.count(".") / max(len(txt), 1) > 0.08:
        score -= 0.2
    return max(0.0, min(1.0, score))


def _repair_mojibake(text: str) -> str:
    """Fix common UTF-8-as-latin1/cp1252 mojibake (e.g. 'Miá»…n' -> 'Miễn')."""
    if not text:
        return text
    # Heuristic: mojibake often contains sequences like Ã, Â, Ä, Å, á» , áº
    markers = ("Ã", "Â", "Ä", "Å", "á»", "áº", "Ã¡", "Ã ", "Ãª", "Æ°", "Æ¡")
    if not any(m in text for m in markers):
        return text
    for enc in ("latin-1", "cp1252"):
        try:
            repaired = text.encode(enc, errors="strict").decode("utf-8")
        except (UnicodeEncodeError, UnicodeDecodeError):
            continue
        # Keep the repair only if it looks better (more Vietnamese letters, fewer weird marks).
        if text_quality(repaired) >= text_quality(text) and (
            "ệ" in repaired or "ễ" in repaired or "ộ" in repaired or "ấ" in repaired
            or repaired.count("Ã") < text.count("Ã")
        ):
            return repaired
    return text


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------

def _from_pdf(data: bytes) -> str:
    """Extract text from a PDF: PyMuPDF → pdfplumber → OCR when layer is missing/garbled."""
    candidates: list[str] = []

    pymupdf_text = _pdf_pymupdf_text(data)
    if pymupdf_text:
        candidates.append(pymupdf_text)

    plumber_text = _pdf_pdfplumber_text(data)
    if plumber_text:
        candidates.append(plumber_text)

    best = max(candidates, key=text_quality, default="")
    best_score = text_quality(best)

    # Good enough text layer → skip OCR (fast path).
    if len(best) >= 40 and best_score >= 0.25:
        return best

    ocr = _ocr_pdf(data)
    if ocr and text_quality(ocr) > best_score:
        return ocr
    return best or ocr


def _pdf_pymupdf_text(data: bytes) -> str:
    """PyMuPDF usually preserves Vietnamese Unicode better than pdfplumber."""
    try:
        import fitz  # PyMuPDF
    except Exception as exc:  # noqa: BLE001
        logger.debug("extract_text: PyMuPDF unavailable: %s", exc)
        return ""
    try:
        parts: list[str] = []
        with fitz.open(stream=data, filetype="pdf") as doc:
            for page in doc:
                # "text" mode keeps reading order; sort=True helps multi-column layouts.
                parts.append(page.get_text("text", sort=True) or "")
        return "\n".join(parts).strip()
    except Exception as exc:  # noqa: BLE001
        logger.warning("extract_text: PyMuPDF text extraction failed: %s", exc)
        return ""


def _pdf_pdfplumber_text(data: bytes) -> str:
    try:
        import pdfplumber
    except Exception as exc:  # noqa: BLE001
        logger.debug("extract_text: pdfplumber unavailable: %s", exc)
        return ""
    try:
        parts: list[str] = []
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            for page in pdf.pages:
                # x_tolerance/y_tolerance reduce character-splitting on Vietnamese fonts.
                parts.append(page.extract_text(x_tolerance=2, y_tolerance=3) or "")
        return "\n".join(parts).strip()
    except Exception as exc:  # noqa: BLE001
        logger.warning("extract_text: PDF text-layer extraction failed: %s", exc)
        return ""


def _ocr_pdf(data: bytes) -> str:
    """OCR a scanned PDF using PyMuPDF (rasterize) + pytesseract with Vietnamese ('vie')."""
    try:
        import fitz  # PyMuPDF
        import pytesseract
        from PIL import Image, ImageFilter, ImageOps
    except Exception as exc:  # noqa: BLE001
        logger.warning("extract_text: OCR libs unavailable (%s); scanned PDF cannot be read.", exc)
        return ""

    cmd = os.getenv("TESSERACT_CMD")
    if not cmd:
        default_win = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
        if os.path.exists(default_win):
            cmd = default_win
    if cmd:
        pytesseract.pytesseract.tesseract_cmd = cmd

    tessdata_dir = os.getenv("TESSERACT_TESSDATA_DIR")
    # PSM 6 = assume a uniform block of text (typical for legal pages).
    config_parts = ["--psm", "6"]
    if tessdata_dir:
        config_parts.extend(["--tessdata-dir", tessdata_dir])
    config = " ".join(config_parts)

    try:
        parts: list[str] = []
        with fitz.open(stream=data, filetype="pdf") as doc:
            for page in doc:
                # 300 DPI is a good balance; higher slows bulk import a lot.
                pix = page.get_pixmap(dpi=300, alpha=False)
                img = Image.open(io.BytesIO(pix.tobytes("png")))
                # Light preprocessing: grayscale + autocontrast improves Vietnamese OCR.
                img = ImageOps.grayscale(img)
                img = ImageOps.autocontrast(img)
                img = img.filter(ImageFilter.SHARPEN)
                try:
                    parts.append(pytesseract.image_to_string(img, lang="vie+eng", config=config))
                except Exception:  # noqa: BLE001
                    try:
                        parts.append(pytesseract.image_to_string(img, lang="vie", config=config))
                    except Exception:  # noqa: BLE001
                        parts.append(pytesseract.image_to_string(img, config=config))
        return "\n".join(parts).strip()
    except Exception as exc:  # noqa: BLE001
        logger.warning("extract_text: OCR failed (Tesseract binary/lang missing?): %s", exc)
        return ""


# ---------------------------------------------------------------------------
# DOCX / DOC / RTF
# ---------------------------------------------------------------------------

def _from_docx(data: bytes) -> str:
    """Extract paragraphs and tables from a modern Word (.docx) file."""
    try:
        import docx
        from docx.document import Document as DocumentType
        from docx.table import Table
        from docx.text.paragraph import Paragraph
    except Exception as exc:  # noqa: BLE001
        logger.warning("extract_text: python-docx unavailable: %s", exc)
        return ""

    try:
        document = docx.Document(io.BytesIO(data))
        parts: list[str] = []

        def walk_blocks(parent: DocumentType) -> None:
            # Iterate body children in document order so tables aren't lost / reordered.
            body = parent.element.body
            for child in body.iterchildren():
                tag = child.tag.split("}")[-1]
                if tag == "p":
                    p = Paragraph(child, parent)
                    t = (p.text or "").strip()
                    if t:
                        parts.append(t)
                elif tag == "tbl":
                    table = Table(child, parent)
                    for row in table.rows:
                        cells = [(c.text or "").strip() for c in row.cells]
                        cells = [c for c in cells if c]
                        if cells:
                            parts.append(" | ".join(cells))

        walk_blocks(document)
        if not parts:
            # Fallback: paragraphs only (older python-docx edge cases).
            parts = [p.text for p in document.paragraphs if (p.text or "").strip()]
            for table in document.tables:
                for row in table.rows:
                    cells = [(c.text or "").strip() for c in row.cells if (c.text or "").strip()]
                    if cells:
                        parts.append(" | ".join(cells))
        return "\n".join(parts).strip()
    except Exception as exc:  # noqa: BLE001
        logger.warning("extract_text: DOCX extraction failed: %s", exc)
        return ""


def _from_doc(data: bytes, filename: str) -> str:
    """Extract text from legacy Word (.doc).

    Order of attempts:
      1) Microsoft Word COM automation (Windows, if Word is installed)
      2) LibreOffice / soffice headless convert → .docx → python-docx
      3) antiword / catdoc if present on PATH
    """
    text = _doc_via_win32com(data)
    if text:
        return text
    text = _doc_via_libreoffice(data, filename)
    if text:
        return text
    text = _doc_via_cli(data, filename)
    if text:
        return text
    logger.warning(
        "extract_text: cannot read legacy .doc (%s). Install LibreOffice or Microsoft Word, "
        "or convert the file to .docx/.pdf before upload.",
        filename,
    )
    return ""


def _doc_via_win32com(data: bytes) -> str:
    try:
        import win32com.client  # type: ignore
        import pythoncom  # type: ignore
    except Exception:
        return ""

    tmp_dir = tempfile.mkdtemp(prefix="legal_doc_")
    src = Path(tmp_dir) / "input.doc"
    dst = Path(tmp_dir) / "output.txt"
    try:
        src.write_bytes(data)
        pythoncom.CoInitialize()
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        try:
            doc = word.Documents.Open(str(src), ReadOnly=True)
            # 2 = wdFormatText
            doc.SaveAs(str(dst), FileFormat=2)
            doc.Close(False)
        finally:
            word.Quit()
            pythoncom.CoUninitialize()
        if dst.exists():
            return _decode(dst.read_bytes())
    except Exception as exc:  # noqa: BLE001
        logger.debug("extract_text: win32com .doc failed: %s", exc)
    finally:
        shutil.rmtree(tmp_dir, ignore_errors=True)
    return ""


def _doc_via_libreoffice(data: bytes, filename: str) -> str:
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        # Common Windows install path
        for candidate in (
            r"C:\Program Files\LibreOffice\program\soffice.exe",
            r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        ):
            if os.path.exists(candidate):
                soffice = candidate
                break
    if not soffice:
        return ""

    tmp_dir = tempfile.mkdtemp(prefix="legal_lo_")
    try:
        src_name = Path(filename).name or "input.doc"
        if not src_name.lower().endswith(".doc"):
            src_name += ".doc"
        src = Path(tmp_dir) / src_name
        src.write_bytes(data)
        # Convert to docx (preserves more structure than plain txt).
        proc = subprocess.run(
            [soffice, "--headless", "--norestore", "--convert-to", "docx", "--outdir", tmp_dir, str(src)],
            capture_output=True,
            timeout=120,
            check=False,
        )
        if proc.returncode != 0:
            logger.debug("extract_text: LibreOffice convert failed: %s", proc.stderr[:300])
            return ""
        converted = src.with_suffix(".docx")
        if not converted.exists():
            # LibreOffice may rewrite the stem; pick any .docx in the temp dir.
            docs = list(Path(tmp_dir).glob("*.docx"))
            if not docs:
                return ""
            converted = docs[0]
        return _from_docx(converted.read_bytes())
    except Exception as exc:  # noqa: BLE001
        logger.debug("extract_text: LibreOffice .doc failed: %s", exc)
        return ""
    finally:
        shutil.rmtree(tmp_dir, ignore_errors=True)


def _doc_via_cli(data: bytes, filename: str) -> str:
    """Last-resort: antiword / catdoc on PATH (common on Linux / WSL)."""
    tool = shutil.which("antiword") or shutil.which("catdoc")
    if not tool:
        return ""
    tmp_dir = tempfile.mkdtemp(prefix="legal_cli_")
    try:
        src = Path(tmp_dir) / (Path(filename).name or "input.doc")
        src.write_bytes(data)
        proc = subprocess.run([tool, str(src)], capture_output=True, timeout=60, check=False)
        if proc.returncode != 0:
            return ""
        # Prefer UTF-8 stdout; fall back to cp1258 / latin-1.
        out = proc.stdout or b""
        return _decode(out)
    except Exception as exc:  # noqa: BLE001
        logger.debug("extract_text: antiword/catdoc failed: %s", exc)
        return ""
    finally:
        shutil.rmtree(tmp_dir, ignore_errors=True)


def _from_rtf(data: bytes) -> str:
    """Minimal RTF stripper (keeps readable text, drops control words)."""
    raw = _decode(data)
    # Remove RTF groups that are usually binary/font tables.
    raw = re.sub(r"\\bin\d+[^}]*", " ", raw)
    raw = re.sub(r"\{\\*\\[^}]+\}", " ", raw)
    # \'hh hex escapes
    def _hex_esc(m: re.Match[str]) -> str:
        try:
            return bytes([int(m.group(1), 16)]).decode("cp1258", errors="ignore")
        except Exception:  # noqa: BLE001
            return ""

    raw = re.sub(r"\\'([0-9a-fA-F]{2})", _hex_esc, raw)
    # Unicode escapes: \uN?
    def _u_esc(m: re.Match[str]) -> str:
        try:
            return chr(int(m.group(1)))
        except Exception:  # noqa: BLE001
            return ""

    raw = re.sub(r"\\u(-?\d+)\??", _u_esc, raw)
    raw = re.sub(r"\\[a-zA-Z]+\d* ?", " ", raw)
    raw = re.sub(r"[{}]", " ", raw)
    return raw


# ---------------------------------------------------------------------------
# Plain / HTML decode
# ---------------------------------------------------------------------------

def _decode(data: bytes) -> str:
    """Decode bytes with encodings common for Vietnamese legal exports."""
    if not data:
        return ""
    # BOM sniffing
    if data.startswith(b"\xff\xfe"):
        try:
            return data.decode("utf-16-le")
        except Exception:  # noqa: BLE001
            pass
    if data.startswith(b"\xfe\xff"):
        try:
            return data.decode("utf-16-be")
        except Exception:  # noqa: BLE001
            pass
    if data.startswith(b"\xef\xbb\xbf"):
        try:
            return data.decode("utf-8-sig")
        except Exception:  # noqa: BLE001
            pass

    # Prefer encodings that keep Vietnamese intact. Score each successful decode.
    best = ""
    best_score = -1.0
    for enc in ("utf-8", "utf-8-sig", "cp1258", "windows-1258", "cp1252", "latin-1", "utf-16"):
        try:
            candidate = data.decode(enc)
        except Exception:  # noqa: BLE001
            continue
        score = text_quality(candidate) if len(candidate) >= 20 else (0.1 if candidate else 0.0)
        # Prefer UTF-8 when scores are close (legal corpora are usually UTF-8).
        if enc.startswith("utf-8"):
            score += 0.05
        if score > best_score:
            best, best_score = candidate, score
    if best:
        return best
    return data.decode("utf-8", errors="replace")


def _strip_html(text: str) -> str:
    text = re.sub(r"(?is)<(script|style)[^>]*>.*?</\1>", " ", text)
    text = re.sub(r"(?s)<[^>]+>", " ", text)
    # HTML entities (minimal set; full decode would need html.unescape)
    import html as html_lib

    text = html_lib.unescape(text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n\s*\n\s*\n+", "\n\n", text)
    return text.strip()
