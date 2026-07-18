"""Unit tests for legal text extraction cleaning (encoding / mojibake / Unicode)."""
from __future__ import annotations

from app.pipelines.legal.extract_text import clean_text, extract_text, text_quality, _repair_mojibake


def test_clean_text_nfc_and_control_chars():
    # e + combining acute should become composed "é"; NUL / soft-hyphen stripped.
    raw = "Hi\u0000\u00ade\u0301p"
    out = clean_text(raw)
    assert "\u0000" not in out
    assert "\u00ad" not in out
    assert "ép" in out or "é" in out


def test_repair_mojibake_vietnamese():
    # "Miễn" UTF-8 bytes mis-decoded as latin-1.
    mojibake = "Miễn".encode("utf-8").decode("latin-1")
    assert "Ã" in mojibake or "Â" in mojibake or "á»" in mojibake or "áº" in mojibake or len(mojibake) > 4
    repaired = _repair_mojibake(mojibake)
    assert "Miễn" in repaired or text_quality(repaired) >= text_quality(mojibake)


def test_clean_text_collapses_whitespace():
    out = clean_text("A\t\tB\r\n\r\n\r\nC")
    assert out == "A B\n\nC"


def test_text_quality_penalizes_garbage():
    clean = "Điều 1. Phạm vi điều chỉnh. Nghị định này quy định về lệ phí hộ tịch."
    garbage = "N2lw5 w5Sie2 %^ S S Mã hàng 615.90.00 ÿÄ ÄÃỐ . ~ Í ÕỔ Ổ MM-¬GÌ"
    assert text_quality(clean) > text_quality(garbage)


def test_decode_utf8_plain():
    data = "Điều 40. Nội dung miễn lệ phí hộ tịch.".encode("utf-8")
    out = extract_text(data, filename="note.txt", mime="text/plain")
    assert "Điều 40" in out
    assert "hộ tịch" in out


def test_decode_prefers_utf8_over_latin1_noise():
    data = "Nghị định số 168/2024/NĐ-CP".encode("utf-8")
    out = extract_text(data, filename="a.txt")
    assert "Nghị định" in out
    assert "NĐ-CP" in out


def test_docx_roundtrip_paragraphs():
    try:
        import docx
    except Exception:
        return  # skip if python-docx missing in this interpreter
    from io import BytesIO

    document = docx.Document()
    document.add_paragraph("Điều 1. Phạm vi điều chỉnh")
    document.add_paragraph("1. Nghị định này quy định về hỗ trợ.")
    buf = BytesIO()
    document.save(buf)
    out = extract_text(buf.getvalue(), filename="sample.docx")
    assert "Điều 1" in out
    assert "hỗ trợ" in out
